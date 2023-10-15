from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
import binascii
import http
import pathlib
import re, argparse
import io, os
import time
import urllib.request
from typing import Optional, cast, Dict
from urllib.parse import urlparse
from html.parser import HTMLParser
from docx.oxml import OxmlElement

import docx, docx.table
from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import Image
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from bs4 import BeautifulSoup


USABLE_HEIGHT = Inches(8.1)
USABLE_WIDTH = Inches(5.8)
DEFAULT_DPI = 72

MAX_IMAGE_SIZE = 50 * 1024 * 1024

RFC_2397_BASE64 = ";base64"

INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5 

DEFAULT_TABLE_STYLE = None

DEFAULT_PARAGRAPH_STYLE = None

def extract_css_style(html):
    css_styles = {}
    soup = BeautifulSoup(html, 'html.parser')
    style_tags = soup.find_all('style')
    
    for style_tag in style_tags:
        css = style_tag.get_text()
        parsed_css = re.findall(r'(\w+)\s*:\s*([^;]+);', css)
        for prop, value in parsed_css:
            css_styles[prop] = value
    
    return css_styles

def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)


def is_url(url):
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])


def make_image(data: Optional[bytes]) -> io.BytesIO:
    image_buffer = None
    if data:
        image_buffer = io.BytesIO(data)
        try:
            Image.from_blob(image_buffer.getbuffer())
        except UnrecognizedImageError:
            image_buffer = None

    if not image_buffer:
        broken_img_path = pathlib.Path(__file__).parent / "image-broken.png"
        image_buffer = io.BytesIO(broken_img_path.read_bytes())

    return image_buffer


def load_external_image(src: str) -> Optional[bytes]:
    data = None
    retry = 3
    while retry and not data:
        try:
            with urllib.request.urlopen(src) as response:
                size = response.getheader("Content-Length")
                if size and int(size) > MAX_IMAGE_SIZE:
                    break
            retry = 0
        except urllib.error.URLError:
            retry -= 1
            if retry:
                time.sleep(1)
        else:
            if len(data) <= MAX_IMAGE_SIZE:
                return data
    return None


def load_inline_image(src: str) -> Optional[bytes]:
    image_data = None
    header_data = src.split(RFC_2397_BASE64 + ",", maxsplit=1)
    if len(header_data) == 2:
        data = header_data[1]
        try:
            image_data = base64.b64decode(data, validate=True)
        except (binascii.Error, ValueError):
            pass
    return image_data


def load_image(src: str) -> io.BytesIO:
    image_bytes = (
        load_inline_image(src) if src.startswith("data:") else load_external_image(src)
    )
    return make_image(image_bytes)


def remove_last_occurence(ls, x):
    ls.pop(len(ls) - ls[::-1].index(x) - 1)


def remove_whitespace(string, leading=False, trailing=False):
    if leading:
        string = re.sub(r'^\s*\n+\s*', '', string)

    if trailing:
        string = re.sub(r'\s*\n+\s*$', '', string)

    string = re.sub(r'\s*\n\s*', ' ', string)
    return re.sub(r'\s+', ' ', string)

def add_watermark(doc, watermark_text, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    watermark = doc.sections[0].footer.paragraphs[0]
    watermark.alignment = alignment

    run = watermark.add_run()
    font = run.font
    font.name = 'Arial'
    font.size = Pt(36)
    font.color.rgb = RGBColor(128, 128, 128) 

    run.text = watermark_text

    doc.save("output.docx")

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

font_names = {
    'code': 'Courier',
    'pre': 'Courier',
}

styles = {
    'LIST_BULLET': 'List Bullet',
    'LIST_NUMBER': 'List Number',
}


class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_row_selectors = [
            'table > tr',
            'table > thead > tr',
            'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = DEFAULT_TABLE_STYLE

        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE
        self.css_styles = {}
        self.document = Document()

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options['fix-html']
        self.document = self.doc
        self.include_tables = True
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0
        self.css_styles = {}

    def copy_settings_from(self, other):
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def get_cell_html(self, soup):
        return ' '.join([str(i) for i in soup.contents])

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style and 'margin-right' in style:
            margin_left = style['margin-left']
            margin_right = style['margin-right']
            if "auto" in margin_left and "auto" in margin_right:
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin_suffix = re.sub(r'[a-z!]+', '', margin)
            if len(margin_suffix) > 0:
                margin = int(float(margin_suffix))
                if units == 'px':
                    self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))

    def apply_table_styles(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif align == 'right':
                self.table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif align == 'justify':
                self.table.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    def add_styles_to_table(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.table.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style and 'margin-right' in style:
            margin_left = style['margin-left']
            margin_right = style['margin-right']
            if "auto" in margin_left and "auto" in margin_right:
                self.table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin_suffix = re.sub(r'[a-z]+', '', margin)
            if len(margin_suffix) > 0:
                margin = int(float(margin_suffix))
                if units == 'px':
                    self.table.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))

    def add_styles_to_run(self, style):
        if 'color' in style:
            if 'rgb' in style['color']:
                color = re.sub(r'[a-z()]+', '', style['color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['color']:
                color = style['color'].lstrip('#')
                colors = tuple(int(color[i:i + 2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
            self.run.font.color.rgb = RGBColor(*colors)

        if 'background-color' in style:
            if 'rgb' in style['background-color']:
                color = color = re.sub(r'[a-z()]+', '', style['background-color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['background-color']:
                color = style['background-color'].lstrip('#')
                colors = tuple(int(color[i:i + 2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
            self.run.font.highlight_color = WD_COLOR.GRAY_25  # TODO: map colors

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def parse_dict_string(self, string, separator=';'):
        new_string = string.replace(" ", '').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def handle_li(self):
        list_depth = len(self.tags['list'])
        if list_depth:
            list_type = self.tags['list'][-1]
        else:
            list_type = 'ul'

        if list_type == 'ol':
            list_style = styles['LIST_NUMBER']
        else:
            list_style = styles['LIST_BULLET']

        self.paragraph = self.doc.add_paragraph(style=list_style)
        self.paragraph.paragraph_format.left_indent = Inches(min((list_depth * LIST_INDENT), MAX_INDENT))
        self.paragraph.paragraph_format.line_spacing =1

    def add_image_to_cell(self, cell, image):
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image)

    def handle_img(self, current_attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = 'img'
            return
        src = current_attrs.get("src")
        src_is_url = is_url(src)
        height_attr = current_attrs.get("height")
        width_attr = current_attrs.get("width")
        height_px = int(height_attr) if height_attr else None
        width_px = int(width_attr) if width_attr else None

        image = load_image(src)
        size = image_size(image, width_px, height_px)
        if image:
            try:
                if isinstance(self.doc, docx.document.Document):
                    self.doc.add_picture(image, **size)
                else:
                    self.add_image_to_cell(self.doc, image)
            except FileNotFoundError:
                image = None
        if not image:
            if src_is_url:
                self.doc.add_paragraph("<image: %s>" % src)
            else:
                self.doc.add_paragraph("<image: %s>" % get_filename_from_url(src))

    def apply_table_styles(self, style):
        if 'font-family' in style:
            font_name = style['font-family']
            for row in self.doc.tables:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = font_name
        if 'border-collapse' in style:
            border_collapse = style['border-collapse']
            if border_collapse == 'collapse':
                for row in self.doc.tables:
                    for cell in row.cells:
                        cell._tc.borders.top.color = RGBColor(0, 0, 0)
                        cell._tc.borders.top.size = Pt(1)
                        cell._tc.borders.left.color = RGBColor(0, 0, 0)
                        cell._tc.borders.left.size = Pt(1)
                        cell._tc.borders.right.color = RGBColor(0, 0, 0)
                        cell._tc.borders.right.size = Pt(1)
                        cell._tc.borders.bottom.color = RGBColor(0, 0, 0)
                        cell._tc.borders.bottom.size = Pt(1)


    def handle_table(self, current_attrs):
        table_soup = self.tables[self.table_no]
        rows, cols = self.get_table_dimensions(table_soup)
        self.table = self.doc.add_table(rows, cols)
        if self.table_style:
            try:
                self.table.style = self.table_style
            except KeyError as e:
                raise ValueError(f"Unable to apply style {self.table_style}.") from e
        
        if 'style' in current_attrs and self.table:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_table(style)
            self.apply_table_styles(style)

        rows = self.get_table_rows(table_soup)
        cell_row = 0
        for row in rows:
            cols = self.get_table_columns(row)
            cell_col = 0
            for col in cols:
                cell_html = self.get_cell_html(col)
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html
                docx_cell = self.table.cell(cell_row, cell_col)
                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html, docx_cell)
                cell_col += 1
            cell_row += 1
        if 'style' in current_attrs and self.table:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_table(style)
            self.apply_table_styles(style)
        self.instances_to_skip = len(table_soup.find_all('table'))
        self.skip_tag = 'table'
        self.skip = True
        self.table = None

    def handle_div(self, current_attrs):
        if 'style' in current_attrs and "page-break-after: always" in current_attrs['style']:
            self.doc.add_page_break()

    def handle_link(self, href, text):
        is_external = href.startswith('http')
        rel_id = self.paragraph.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        )

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)

        subrun = self.paragraph.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), "0000EE")
        rPr.append(c)

        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

        subrun._r.append(rPr)
        subrun._r.text = text

        hyperlink.append(subrun._r)

        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        if tag == 'head':
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return

            return
        current_attrs = dict(attrs)

        if tag == 'span':
            self.tags['span'].append(current_attrs)
            return
        elif tag == 'ol' or tag == 'ul':
            self.tags['list'].append(tag)
            return  # don't apply styles for now
        elif tag == 'br':
            self.run.add_break()
            return
        if tag == 'style':
            self.css_styles.update(extract_css_style(self.get_text()))
            return
        self.tags[tag] = current_attrs
        if tag in ['p', 'pre']:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        elif tag == 'li':
            self.handle_li()

        elif tag == "hr":

            self.paragraph = self.doc.add_paragraph()
            pPr = self.paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pPr.insert_element_before(pBdr,
                                      'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                      'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                      'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                      'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                      'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                      'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                                      'w:pPrChange'
                                      )
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)

        elif re.match('h[1-9]', tag):
            if isinstance(self.doc, docx.document.Document):
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()

        elif tag == 'img':
            self.handle_img(current_attrs)
            self.paragraph = self.doc.paragraphs[-1]

        elif tag == 'table':
            self.handle_table(current_attrs)
            return

        elif tag == "div":
            self.handle_div(current_attrs)

        if tag in ['p', 'li', 'pre']:
            self.run = self.paragraph.add_run()

        if not self.include_styles:
            return
        if 'style' in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_paragraph(style)

    def handle_endtag(self, tag):
        if self.skip:
            if not tag == self.skip_tag:
                return

            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return

            self.skip = False
            self.skip_tag = None
            self.paragraph = None

        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag == 'ol' or tag == 'ul':
            remove_last_occurence(self.tags['list'], tag)
            return
        elif tag == 'table':
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None

        if tag in self.tags:
            self.tags.pop(tag)

    def handle_data(self, data):
        if self.skip:
            return
        if 'pre' not in self.tags:
            data = remove_whitespace(data, True, True)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            self.run = self.paragraph.add_run(data)
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)

            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name

    def ignore_nested_tables(self, tables_soup):
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_table_rows(self, table_soup):
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def get_table_columns(self, row):
        return row.find_all(['th', 'td'], recursive=False) if row else []

    def get_table_dimensions(self, table_soup):
        rows = self.get_table_rows(table_soup)
        cols = self.get_table_columns(rows[0]) if rows else []
        return len(rows), len(cols)

    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = str(self.soup)
        
        if self.include_tables:

            self.get_tables()

        self.process_watermark_tags()
        self.feed(html)

    def process_watermark_tags(self):
        if not self.soup:
            return
        div_tags = self.soup.find_all('div')
        for div_tag in div_tags:
            if 'id' in div_tag.attrs:
                ids = div_tag['id']
                if 'watermark' in ids:
                    watermark_text = div_tag.get_text()
                    print(watermark_text)
                    add_watermark(self.doc,watermark_text)
    

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    def parse_html_file(self, filename_html, filename_docx=None):
        with open(filename_html, 'r') as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = '%s/new_docx_file_%s' % (path, filename)
        self.doc.save('%s.docx' % filename_docx)

    def parse_html_string(self, html, filename_docx=None):
        self.set_initial_attrs()
        self.run_process(html)
        
        if not filename_docx:
            filename_docx = '/Users/pranavaggarwal/Documents/stgi/static/output'  # Default filename
        
        self.doc.save('/Users/pranavaggarwal/Documents/stgi/static/output.docx')
        return filename_docx


def image_size(
    image_buffer: io.BytesIO,
    width_px: Optional[int] = None,
    height_px: Optional[int] = None,
) -> Dict[str, int]:
    image = Image.from_blob(image_buffer.getbuffer())

    if height_px is None:
        height = image.px_height / image.vert_dpi
    else:
        height = height_px / DEFAULT_DPI
    if width_px is None:
        width = image.px_width / image.horz_dpi
    else:
        width = width_px / DEFAULT_DPI

    height = Inches(height)
    width = Inches(width)

    size = {}
    if width > USABLE_WIDTH:
        new_height = round(image.px_height / (image.px_width / USABLE_WIDTH))
        if new_height > USABLE_HEIGHT:
            size["height"] = USABLE_HEIGHT
        else:
            size["width"] = USABLE_WIDTH
    elif height > USABLE_HEIGHT:
        new_width = round(image.px_width / (image.px_height / USABLE_HEIGHT))
        if new_width > USABLE_WIDTH:
            size["width"] = USABLE_WIDTH
        else:
            size["height"] = USABLE_HEIGHT
    else:
        if width_px is not None:
            size["width"] = width
        if height_px is not None:
            size["height"] = height
    return size